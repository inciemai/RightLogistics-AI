from flask import Flask, request, jsonify
import os
import base64
import re
from datetime import datetime
import google.generativeai as genai
from dotenv import load_dotenv
import logging
import io
from PIL import Image
import fitz  # PyMuPDF for PDF processing
from docx import Document
import openpyxl
from openpyxl.drawing.image import Image as OpenpyxlImage
import tempfile
import json

logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY not found in environment variables")

genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-1.5-flash')

app = Flask(__name__)

def enhance_image_for_ocr(image_data):
    """Enhance image quality for better OCR results"""
    try:
        # Convert base64 to PIL Image
        if isinstance(image_data, str):
            image_bytes = base64.b64decode(image_data)
        else:
            image_bytes = image_data
            
        img = Image.open(io.BytesIO(image_bytes))
        
        # Convert to RGB if needed
        if img.mode in ('RGBA', 'LA'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'RGBA':
                background.paste(img, mask=img.split()[-1])
            else:
                background.paste(img)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Enhance image quality
        width, height = img.size
        
        # Resize if image is too small (improve OCR accuracy)
        if width < 800 or height < 600:
            scale_factor = max(800/width, 600/height)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Save enhanced image
        output_buffer = io.BytesIO()
        img.save(output_buffer, format='PNG', quality=95)
        enhanced_data = output_buffer.getvalue()
        
        return base64.b64encode(enhanced_data).decode('utf-8')
    except Exception as e:
        logger.warning(f"Image enhancement failed, using original: {str(e)}")
        return image_data if isinstance(image_data, str) else base64.b64encode(image_data).decode('utf-8')

def encode_image_to_base64(image_data):
    """Encode image data to base64 string with enhancement"""
    if isinstance(image_data, bytes):
        enhanced_data = enhance_image_for_ocr(image_data)
        return enhanced_data
    else:
        enhanced_data = enhance_image_for_ocr(image_data.read())
        return enhanced_data

def extract_text_from_word(file_stream):
    """Extract text from Word document with better formatting"""
    try:
        doc = Document(file_stream)
        text_content = []
        
        # Extract paragraphs with better formatting
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Preserve some formatting cues
                text = paragraph.text.strip()
                if paragraph.style.name.startswith('Heading'):
                    text = f"[HEADER] {text}"
                text_content.append(text)
        
        # Extract tables with improved structure
        for table in doc.tables:
            table_content = ["[TABLE START]"]
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    if cell_text:
                        row_text.append(cell_text)
                    else:
                        row_text.append("")
                if row_text and any(cell for cell in row_text):
                    if row_idx == 0:  # Likely header row
                        table_content.append(f"[HEADER ROW] {' | '.join(row_text)}")
                    else:
                        table_content.append(f"[DATA ROW] {' | '.join(row_text)}")
            table_content.append("[TABLE END]")
            text_content.extend(table_content)
        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        raise Exception(f"Failed to extract text from Word document: {str(e)}")

def extract_data_from_excel_with_formulas(file_stream):
    """Extract data from Excel file with improved structure recognition"""
    try:
        workbook_formulas = openpyxl.load_workbook(file_stream, data_only=False)
        file_stream.seek(0)
        workbook_values = openpyxl.load_workbook(file_stream, data_only=True)
        
        text_content = []
        
        for sheet_name in workbook_formulas.sheetnames:
            sheet_formulas = workbook_formulas[sheet_name]
            sheet_values = workbook_values[sheet_name]
            sheet_content = [f"[SHEET: {sheet_name}]"]
            
            max_row = sheet_formulas.max_row
            max_col = sheet_formulas.max_column
            
            # Try to identify header rows and data structure
            header_detected = False
            
            for row in range(1, max_row + 1):
                row_data = []
                has_content = False
                
                for col in range(1, max_col + 1):
                    cell_formula = sheet_formulas.cell(row, col)
                    cell_value = sheet_values.cell(row, col)
                    
                    cell_content = ""
                    if cell_formula.value is not None:
                        if isinstance(cell_formula.value, str) and cell_formula.value.startswith('='):
                            calculated_value = cell_value.value if cell_value.value is not None else 0
                            cell_content = f"{cell_formula.value}[={calculated_value}]"
                        else:
                            cell_content = str(cell_formula.value)
                        has_content = True
                    elif cell_value.value is not None:
                        cell_content = str(cell_value.value)
                        has_content = True
                    
                    row_data.append(cell_content)
                
                if has_content:
                    # Detect if this might be a header row
                    if not header_detected and row <= 3:  # Check first 3 rows for headers
                        text_values = [cell for cell in row_data if cell and not cell.replace('.', '').replace('-', '').isdigit()]
                        if len(text_values) >= len([cell for cell in row_data if cell]) * 0.7:  # 70% text content
                            sheet_content.append(f"[HEADER ROW] {' | '.join(row_data)}")
                            header_detected = True
                            continue
                    
                    sheet_content.append(f"[DATA ROW] {' | '.join(row_data)}")
            
            if len(sheet_content) > 1:
                text_content.extend(sheet_content)
        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error extracting data from Excel file: {str(e)}")
        raise Exception(f"Failed to extract data from Excel file: {str(e)}")

def extract_images_from_pdf(file_stream):
    """Extract images from PDF with better quality"""
    try:
        pdf_document = fitz.open(stream=file_stream.read(), filetype="pdf")
        images = []
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(pdf_document, xref)
                
                # Skip very small images (likely decorative)
                if pix.width < 100 or pix.height < 100:
                    continue
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    images.append(img_data)
                else:  # CMYK: convert to RGB first
                    pix1 = fitz.Pixmap(fitz.csRGB, pix)
                    img_data = pix1.tobytes("png")
                    images.append(img_data)
                    pix1 = None
                
                pix = None
        
        pdf_document.close()
        return images
    except Exception as e:
        logger.error(f"Error extracting images from PDF: {str(e)}")
        raise Exception(f"Failed to extract images from PDF: {str(e)}")

def extract_text_from_pdf(file_stream):
    """Extract text from PDF with multiple methods and comprehensive debugging"""
    try:
        pdf_document = fitz.open(stream=file_stream.read(), filetype="pdf")
        all_text_methods = []
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            
            # Method 1: Simple text extraction
            simple_text = page.get_text()
            
            # Method 2: Text with blocks
            blocks_text = page.get_text("blocks")
            blocks_content = []
            for block in blocks_text:
                if len(block) >= 5 and block[4].strip():  # block[4] contains text
                    blocks_content.append(block[4].strip())
            
            # Method 3: Dictionary method with detailed structure
            dict_text = page.get_text("dict")
            structured_content = []
            for block in dict_text["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            if span["text"].strip():
                                line_text += span["text"] + " "
                        if line_text.strip():
                            structured_content.append(line_text.strip())
            
            # Method 4: HTML extraction (often captures more details)
            try:
                html_text = page.get_text("html")
                # Extract text from HTML tags
                import re
                html_clean = re.sub(r'<[^>]+>', ' ', html_text)
                html_clean = ' '.join(html_clean.split())
            except:
                html_clean = ""
            
            # Combine all methods
            page_content = [f"\n=== PAGE {page_num + 1} - METHOD 1 (SIMPLE) ==="]
            page_content.append(simple_text)
            
            page_content.append(f"\n=== PAGE {page_num + 1} - METHOD 2 (BLOCKS) ===")
            page_content.extend(blocks_content)
            
            page_content.append(f"\n=== PAGE {page_num + 1} - METHOD 3 (STRUCTURED) ===")
            page_content.extend(structured_content)
            
            if html_clean:
                page_content.append(f"\n=== PAGE {page_num + 1} - METHOD 4 (HTML) ===")
                page_content.append(html_clean)
            
            all_text_methods.extend(page_content)
        
        pdf_document.close()
        final_text = "\n".join(all_text_methods)
        
        # Enhanced logging
        logger.info(f"PDF extraction completed. Total text length: {len(final_text)}")
        logger.info(f"First 2000 characters of extracted text:\n{final_text[:2000]}")
        
        # Check if we actually extracted meaningful content
        meaningful_words = ['invoice', 'total', 'date', 'amount', '₹', '$', '@', '+91']
        found_keywords = [word for word in meaningful_words if word.lower() in final_text.lower()]
        logger.info(f"Keywords found in extracted text: {found_keywords}")
        
        if len(final_text.strip()) < 50:
            logger.warning("Very little text extracted from PDF - may need OCR processing")
            
        return final_text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def process_file_content(file_stream, filename):
    """Enhanced file processing with better PDF handling"""
    file_extension = filename.rsplit('.', 1)[1].lower()
    
    try:
        if file_extension in ['png', 'jpg', 'jpeg', 'webp']:
            image_data = encode_image_to_base64(file_stream)
            return {"type": "image", "data": image_data, "extension": file_extension}
        
        elif file_extension == 'pdf':
            file_stream.seek(0)
            
            # First, try to extract images for OCR
            pdf_images = extract_images_from_pdf(file_stream)
            
            # Always try text extraction first
            file_stream.seek(0)
            text_content = extract_text_from_pdf(file_stream)
            
            # Check if text extraction was successful
            meaningful_content = len(text_content.strip()) > 100
            has_invoice_keywords = any(keyword in text_content.lower() for keyword in ['invoice', 'total', 'amount', 'date'])
            
            logger.info(f"PDF text extraction - Length: {len(text_content)}, Has keywords: {has_invoice_keywords}")
            
            if meaningful_content and has_invoice_keywords:
                # Use text extraction
                return {"type": "text", "data": text_content, "extension": file_extension}
            elif pdf_images:
                # Fallback to OCR if text extraction failed but images exist
                logger.info("Text extraction insufficient, falling back to OCR")
                best_image = max(pdf_images, key=len)
                image_data = encode_image_to_base64(best_image)
                return {"type": "image", "data": image_data, "extension": file_extension}
            else:
                # Use whatever text we got
                logger.warning("Limited text extracted and no images found, proceeding with available text")
                return {"type": "text", "data": text_content, "extension": file_extension}
        
        elif file_extension in ['doc', 'docx']:
            text_content = extract_text_from_word(file_stream)
            return {"type": "text", "data": text_content, "extension": file_extension}
        
        elif file_extension in ['xls', 'xlsx']:
            text_content = extract_data_from_excel_with_formulas(file_stream)
            return {"type": "text", "data": text_content, "extension": file_extension}
        
        else:
            raise Exception("Unsupported file type")
            
    except Exception as e:
        raise Exception(f"Error processing file: {str(e)}")

def validate_and_clean_extracted_data(extracted_data):
    """Enhanced validation and cleaning of extracted data"""
    cleaned_data = extracted_data.copy()
    
    # Clean and validate date
    if cleaned_data.get('date'):
        date_str = str(cleaned_data['date']).strip()
        try:
            # Handle various date formats including DD/MM/YYYY
            for fmt in ['%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y', '%m-%d-%Y', '%Y/%m/%d']:
                try:
                    parsed_date = datetime.strptime(date_str, fmt)
                    cleaned_data['date'] = parsed_date.strftime('%Y-%m-%d')
                    break
                except ValueError:
                    continue
        except:
            pass  # Keep original if parsing fails
    
    # Clean numeric fields and preserve currency symbols in context
    numeric_fields = ['subtotal', 'tax', 'total']
    for field in numeric_fields:
        if cleaned_data.get(field):
            value = str(cleaned_data[field])
            # Extract numeric value but keep currency info
            currency_match = re.search(r'[₹$€£¥]', value)
            currency = currency_match.group() if currency_match else ''
            
            # Clean numeric value
            numeric_value = re.sub(r'[^\d.,\-]', '', value)
            numeric_value = numeric_value.replace(',', '')
            
            try:
                float(numeric_value)
                if currency:
                    cleaned_data[field] = f"{currency}{numeric_value}"
                else:
                    cleaned_data[field] = numeric_value
            except ValueError:
                pass  # Keep original if not numeric
    
    # Enhanced item cleaning
    if cleaned_data.get('items') and isinstance(cleaned_data['items'], list):
        cleaned_items = []
        for item in cleaned_data['items']:
            if isinstance(item, dict):
                cleaned_item = {}
                for key, value in item.items():
                    if value and str(value).strip() and str(value).strip().lower() not in ['null', 'none', '', 'n/a']:
                        if key in ['quantity', 'unit_price', 'amount']:
                            # Clean numeric values in items
                            str_value = str(value)
                            currency_match = re.search(r'[₹$€£¥]', str_value)
                            currency = currency_match.group() if currency_match else ''
                            
                            cleaned_value = re.sub(r'[^\d.,\-]', '', str_value)
                            cleaned_value = cleaned_value.replace(',', '')
                            
                            try:
                                float(cleaned_value)
                                if currency and key in ['unit_price', 'amount']:
                                    cleaned_item[key] = f"{currency}{cleaned_value}"
                                else:
                                    cleaned_item[key] = cleaned_value
                            except ValueError:
                                cleaned_item[key] = str(value).strip()
                        else:
                            cleaned_item[key] = str(value).strip()
                
                # Add item if it has meaningful content
                if cleaned_item.get('description') or cleaned_item.get('amount'):
                    cleaned_items.append(cleaned_item)
        
        cleaned_data['items'] = cleaned_items
    
    return cleaned_data

def get_extraction_prompt(file_extension):
    """Get enhanced extraction prompt with better instructions and examples"""
    
    base_prompt = """
You are an expert invoice data extraction system. You MUST extract ALL available information from the provided invoice document.

CRITICAL EXTRACTION RULES:
1. NEVER return null for fields that have visible data in the document
2. Look for information in ALL parts of the document - headers, body, footer, contact sections
3. Extract EXACT values as they appear, including currency symbols and formatting
4. Be thorough - check every line and section for relevant data

REQUIRED FIELDS TO EXTRACT:

1. INVOICE NUMBER: Look for patterns like "Invoice No:", "INV", "Invoice #", "Bill No", etc.
2. DATE: Look for "Date:", "Invoice Date:", or any date in DD/MM/YYYY, MM/DD/YYYY, YYYY-MM-DD format
3. CUSTOMER NAME: Look in "INVOICE TO:", "Bill To:", "Customer:", or contact information sections
4. CUSTOMER CONTACT: Extract phone numbers (+91, mobile) and email addresses from anywhere in document
5. CUSTOMER ADDRESS: Complete address from "INVOICE TO:" or customer information section
6. ITEMS/SERVICES: Look for:
   - Description of services/products
   - Quantities (if applicable, otherwise use "1")
   - Unit prices or total amounts
   - Line items in tables or lists

7. FINANCIAL TOTALS:
   - Subtotal (amount before tax)
   - Tax amount (GST, VAT, Tax, etc.)
   - Total amount (final payable amount)

SEARCH STRATEGY:
- Scan the ENTIRE document systematically
- Check headers, footers, and all text sections
- Look for currency symbols (₹, $, etc.) to identify amounts
- Phone numbers often start with +91 (India), look for mobile/contact info
- Email addresses contain @ symbol
- Addresses typically have city, state, postal codes

AMOUNT EXTRACTION:
- Always extract the numerical value with currency symbol if present
- For single-item invoices, use the total amount as both unit price and amount
- If no separate tax is shown, tax can be null but subtotal and total should be extracted
"""

    if file_extension == 'pdf':
        specific_instructions = """
PDF-SPECIFIC INSTRUCTIONS:
- The document contains both simple text and structured text
- Look in BOTH "SIMPLE TEXT:" and "STRUCTURED TEXT:" sections
- Headers are marked with [HEADER] and [SUBHEADER]
- Pay attention to layout and positioning
- Contact information might be in separate sections
- Tables might be represented as aligned text
"""
    elif file_extension in ['xls', 'xlsx']:
        specific_instructions = """
EXCEL-SPECIFIC INSTRUCTIONS:
- Look for [HEADER ROW] and [DATA ROW] markers
- Extract values from formula results [=value]
- Handle multiple sheets if present
- Table structure is clearly marked
"""
    else:
        specific_instructions = """
DOCUMENT-SPECIFIC INSTRUCTIONS:
- Use structural markers like [HEADER], [TABLE START/END]
- Handle formatted text and tables appropriately
- Look for visual formatting cues
"""

    json_format = """
RESPONSE FORMAT - Return ONLY valid JSON:
{
    "invoice_number": "extracted_invoice_number_or_null",
    "date": "extracted_date_or_null", 
    "customer_name": "extracted_customer_name_or_null",
    "customer_contact": "extracted_phone_and_email_or_null",
    "customer_address": "complete_extracted_address_or_null",
    "items": [
        {
            "quantity": "1_or_extracted_quantity",
            "description": "extracted_service_or_product_description",
            "unit_price": "extracted_price_or_amount",
            "amount": "extracted_total_amount_for_item"
        }
    ],
    "subtotal": "extracted_subtotal_or_null",
    "tax": "extracted_tax_amount_or_null", 
    "total": "extracted_total_amount"
}

VALIDATION CHECKLIST before responding:
✓ Did I check ALL text sections for customer information?
✓ Did I find the invoice number and date?
✓ Did I extract the service/product description?
✓ Did I find contact information (phone/email)?
✓ Did I extract the complete address?
✓ Did I identify all monetary amounts?
✓ Are there any fields I marked as null that actually have data?

CRITICAL: If you see ANY data in the document, you MUST extract it. Do NOT return null for visible information.
"""

    return base_prompt + specific_instructions + json_format

def calculate_item_amounts_excel_only(items, file_extension):
    """Calculate item amounts only for Excel files with better validation"""
    if file_extension not in ['xls', 'xlsx']:
        return items
    
    calculated_items = []
    
    for item in items:
        calculated_item = item.copy()
        
        try:
            quantity_str = str(item.get('quantity', '0')).replace(',', '')
            unit_price_str = str(item.get('unit_price', '0')).replace(',', '')
            
            # Extract numeric values more carefully
            quantity = 0
            unit_price = 0
            
            # Handle formula results
            if '[=' in quantity_str:
                match = re.search(r'\[=([^\]]+)\]', quantity_str)
                if match:
                    quantity = float(match.group(1))
                else:
                    quantity = float(re.sub(r'[^\d.-]', '', quantity_str))
            else:
                quantity = float(re.sub(r'[^\d.-]', '', quantity_str)) if quantity_str else 0
            
            if '[=' in unit_price_str:
                match = re.search(r'\[=([^\]]+)\]', unit_price_str)
                if match:
                    unit_price = float(match.group(1))
                else:
                    unit_price = float(re.sub(r'[^\d.-]', '', unit_price_str))
            else:
                unit_price = float(re.sub(r'[^\d.-]', '', unit_price_str)) if unit_price_str else 0
            
            # Calculate amount if missing or formula-based
            amount = item.get('amount', '')
            if not amount or '=' in str(amount) or '[=' in str(amount):
                calculated_amount = round(quantity * unit_price, 2)
                calculated_item['amount'] = str(calculated_amount)
            else:
                # Clean existing amount
                amount_str = str(amount).replace(',', '')
                if '[=' in amount_str:
                    match = re.search(r'\[=([^\]]+)\]', amount_str)
                    if match:
                        calculated_item['amount'] = str(float(match.group(1)))
                else:
                    try:
                        calculated_item['amount'] = str(float(re.sub(r'[^\d.-]', '', amount_str)))
                    except:
                        calculated_item['amount'] = str(amount)
                        
        except (ValueError, TypeError) as e:
            logger.warning(f"Error calculating item amount: {e}")
            pass
            
        calculated_items.append(calculated_item)
    
    return calculated_items

def calculate_totals_excel_only(items, extracted_data, file_extension):
    """Calculate totals only for Excel files with better validation"""
    if file_extension not in ['xls', 'xlsx']:
        return {
            'subtotal': extracted_data.get('subtotal'),
            'tax': extracted_data.get('tax'),
            'total': extracted_data.get('total')
        }
    
    try:
        # Calculate subtotal from items
        subtotal = 0
        for item in items:
            try:
                amount_str = str(item.get('amount', '0')).replace(',', '')
                amount = float(re.sub(r'[^\d.-]', '', amount_str))
                subtotal += amount
            except (ValueError, TypeError):
                pass
        
        # Handle tax with formula support
        tax_value = 0
        tax_field = extracted_data.get('tax')
        if tax_field and str(tax_field).strip() not in ['null', 'None', '']:
            tax_str = str(tax_field)
            try:
                if '[=' in tax_str:
                    match = re.search(r'\[=([^\]]+)\]', tax_str)
                    if match:
                        tax_value = float(match.group(1))
                else:
                    tax_value = float(re.sub(r'[^\d.-]', '', tax_str))
            except (ValueError, TypeError):
                pass
        
        # Calculate total
        total = round(subtotal + tax_value, 2)
        
        return {
            'subtotal': str(round(subtotal, 2)),
            'tax': str(round(tax_value, 2)) if tax_value > 0 else extracted_data.get('tax'),
            'total': str(total)
        }
        
    except Exception as e:
        logger.error(f"Error calculating totals: {str(e)}")
        return {
            'subtotal': extracted_data.get('subtotal'),
            'tax': extracted_data.get('tax'),
            'total': extracted_data.get('total')
        }

def extract_invoice_data_with_gemini(content_data):
    """Use Google Gemini AI to extract invoice data with enhanced prompting and debugging"""
    try:
        file_extension = content_data.get('extension', 'unknown')
        prompt = get_extraction_prompt(file_extension)
        
        logger.info(f"Processing {file_extension} file with Gemini AI")
        
        if content_data["type"] == "image":
            # For images, use enhanced prompt
            image_prompt = prompt + "\n\nANALYZE THIS INVOICE IMAGE CAREFULLY. Extract ALL visible text and data."
            response = model.generate_content([image_prompt, {"mime_type": "image/png", "data": content_data["data"]}])
        elif content_data["type"] == "text":
            # Log the text content for debugging
            text_content = content_data['data']
            logger.info(f"Text content length: {len(text_content)}")
            logger.info(f"Text preview: {text_content[:500]}...")
            
            full_prompt = f"""{prompt}

DOCUMENT CONTENT TO ANALYZE:
{text_content}

IMPORTANT: The above document content contains an invoice. Extract ALL the information you can see. Do not return null for any field that has visible data."""
            
            response = model.generate_content(full_prompt)
        else:
            raise Exception("Invalid content type")
        
        response_text = response.text
        logger.info(f"AI Response length: {len(response_text)}")
        logger.info(f"AI Response preview: {response_text[:300]}...")
        
        # Clean and extract JSON from response
        json_str = response_text.strip()
        
        # Remove markdown code blocks if present
        if json_str.startswith('```json'):
            json_str = json_str.replace('```json', '').replace('```', '').strip()
        elif json_str.startswith('```'):
            json_str = json_str.replace('```', '').strip()
        
        # Try to find JSON in the response
        json_match = re.search(r'(\{[\s\S]*\})', json_str)
        if json_match:
            json_str = json_match.group(1)
        
        try:
            extracted_data = json.loads(json_str)
            logger.info(f"Successfully parsed JSON: {extracted_data}")
            
            # Validate and clean the data
            extracted_data = validate_and_clean_extracted_data(extracted_data)
            
            # Post-process calculations only for Excel files
            if file_extension in ['xls', 'xlsx'] and 'items' in extracted_data and extracted_data['items']:
                calculated_items = calculate_item_amounts_excel_only(extracted_data['items'], file_extension)
                extracted_data['items'] = calculated_items
                
                calculated_totals = calculate_totals_excel_only(calculated_items, extracted_data, file_extension)
                extracted_data.update(calculated_totals)
            
            extracted_data['file_extension'] = file_extension
            
            # Log final extracted data
            logger.info(f"Final extracted data: {extracted_data}")
            
            return extracted_data
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {e}")
            logger.error(f"Problematic JSON string: {json_str}")
            
            # Try to fix common JSON issues
            fixed_json = json_str.replace("'", '"').replace('True', 'true').replace('False', 'false').replace('None', 'null')
            
            try:
                extracted_data = json.loads(fixed_json)
                logger.info("Successfully parsed fixed JSON")
                return extracted_data
            except:
                raise Exception(f"Failed to parse AI response as JSON. Response was: {response_text[:500]}...")
                
    except Exception as e:
        logger.error(f"Error in Gemini API processing: {str(e)}")
        logger.error(f"Content data type: {content_data.get('type')}")
        logger.error(f"File extension: {content_data.get('extension')}")
        raise Exception(f"AI processing error: {str(e)}")

def create_api_response(status, message, data=None):
    """Create standardized API response"""
    response = {
        "status": status,
        "message": message,
        "data": data
    }
    return response

@app.route('/extract-invoice', methods=['POST'])
def extract_invoice():
    """API endpoint to extract data from invoice files with enhanced processing"""
    
    try:
        if 'file' not in request.files:
            response = create_api_response("error", "No file provided", None)
            return jsonify(response), 400
        
        uploaded_file = request.files['file']
        
        if uploaded_file.filename == '':
            response = create_api_response("error", "Empty filename", None)
            return jsonify(response), 400

        allowed_extensions = {'png', 'jpg', 'jpeg', 'webp', 'pdf', 'doc', 'docx', 'xls', 'xlsx'}
        file_extension = uploaded_file.filename.rsplit('.', 1)[1].lower() if '.' in uploaded_file.filename else ''
        
        if not (file_extension in allowed_extensions):
            response = create_api_response(
                "error", 
                "File type not allowed. Please upload PNG, JPG, JPEG, WEBP, PDF, DOC, DOCX, XLS, or XLSX files", 
                None
            )
            return jsonify(response), 400

        file_stream = io.BytesIO(uploaded_file.read())

        try:
            content_data = process_file_content(file_stream, uploaded_file.filename)
            logger.info(f"Successfully processed {file_extension} file")
        except Exception as e:
            logger.error(f"File processing error: {str(e)}")
            response = create_api_response("error", str(e), None)
            return jsonify(response), 400

        try:
            extracted_data = extract_invoice_data_with_gemini(content_data)

            logger.info(f"Extraction completed for {file_extension} file")
            logger.info(f"Items extracted: {len(extracted_data.get('items', []))}")
            

            if file_extension in ['xls', 'xlsx']:
                message = f"Invoice data extracted successfully from {file_extension.upper()} file with enhanced calculation processing"
            else:
                message = f"Invoice data extracted successfully from {file_extension.upper()} file with enhanced OCR and structure recognition"
            
            response = create_api_response("success", message, extracted_data)
            return jsonify(response), 200
            
        except Exception as e:
            logger.error(f"Extraction error: {str(e)}")
            response = create_api_response("error", str(e), None)
            return jsonify(response), 500
    
    except Exception as e:
        logger.error(f"Unexpected error processing invoice: {str(e)}")
        response = create_api_response("error", f"Unexpected error: {str(e)}", None)
        return jsonify(response), 500

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    response = create_api_response("error", "Endpoint not found", None)
    return jsonify(response), 404

@app.errorhandler(405)
def method_not_allowed(error):
    """Handle 405 errors"""
    response = create_api_response("error", "Method not allowed", None)
    return jsonify(response), 405

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors"""
    response = create_api_response("error", "Internal server error", None)
    return jsonify(response), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port, debug=False)